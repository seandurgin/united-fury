#!/usr/bin/env python3
import os, sqlite3, logging, asyncio, httpx, base64, json, re, requests, msal, io
from datetime import datetime
from email.mime.text import MIMEText
import anthropic
from telegram import Update
from telegram.ext import Application, MessageHandler, CommandHandler, filters, ContextTypes
from telegram.constants import ChatAction
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build

logging.basicConfig(format="%(asctime)s [%(levelname)s] %(name)s: %(message)s", level=logging.INFO,
    handlers=[logging.StreamHandler(), logging.FileHandler("/var/log/clawdia.log", encoding="utf-8")])
log = logging.getLogger("clawdia")

# Plaid finance module (optional — only loads if configured)
try:
    from plaid_finance import get_accounts, get_transactions, get_debt_snapshot, spending_by_category, exchange_public_token
    PLAID_ENABLED = True
except Exception as _pe:
    PLAID_ENABLED = False
    log.warning("Plaid not available: %s", _pe)

# Plaid finance module (optional — only loads if configured)
try:
    from plaid_finance import get_accounts, get_transactions, get_debt_snapshot, spending_by_category, exchange_public_token
    PLAID_ENABLED = True
except Exception as _pe:
    PLAID_ENABLED = False
    log.warning("Plaid not available: %s", _pe)

TELEGRAM_TOKEN    = os.environ["TELEGRAM_TOKEN"]
ANTHROPIC_KEY     = os.environ["ANTHROPIC_API_KEY"]
BRAVE_KEY         = os.environ.get("BRAVE_API_KEY", "")
OWNER_TELEGRAM_ID = int(os.environ.get("OWNER_TELEGRAM_ID", "0"))
DB_PATH           = os.environ.get("DB_PATH", "/var/lib/clawdia/memory.db")
GOOGLE_TOKEN      = "/etc/clawdia/google_token.json"
FAMILY_TOKEN      = "/etc/clawdia/google_token_family.json"
MS_TOKEN          = "/etc/clawdia/ms_token.json"
MODEL             = "claude-sonnet-4-6"
MAX_HISTORY       = 40
MAX_MEMORY_CHARS  = 8000
GOOGLE_SCOPES     = ['https://www.googleapis.com/auth/gmail.modify','https://www.googleapis.com/auth/calendar','https://www.googleapis.com/auth/drive','https://www.googleapis.com/auth/contacts.readonly']
MS_SCOPES         = ["Notes.ReadWrite","Mail.Read","Calendars.Read","User.Read"]
MS_CLIENT_ID      = "10fd6347-d39f-40cd-bbff-51a8c2af8471"
MS_AUTHORITY      = "https://login.microsoftonline.com/consumers"
GRAPH_BASE        = "https://graph.microsoft.com/v1.0"

def init_db():
    os.makedirs(os.path.dirname(DB_PATH), exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.executescript("""
        CREATE TABLE IF NOT EXISTS memory (id INTEGER PRIMARY KEY AUTOINCREMENT, category TEXT NOT NULL, key TEXT NOT NULL, value TEXT NOT NULL, created TEXT NOT NULL, updated TEXT NOT NULL, UNIQUE(category, key));
        CREATE TABLE IF NOT EXISTS history (id INTEGER PRIMARY KEY AUTOINCREMENT, chat_id INTEGER NOT NULL, role TEXT NOT NULL, content TEXT NOT NULL, ts TEXT NOT NULL);
        CREATE INDEX IF NOT EXISTS idx_history_chat ON history(chat_id, id);
    """)
    conn.commit(); conn.close()

def get_conn(): return sqlite3.connect(DB_PATH)

def memory_save(category, key, value):
    if not category or not key or not value: return
    now = datetime.utcnow().isoformat()
    with get_conn() as conn:
        conn.execute("INSERT INTO memory(category,key,value,created,updated) VALUES(?,?,?,?,?) ON CONFLICT(category,key) DO UPDATE SET value=excluded.value,updated=excluded.updated",
            (str(category).strip(), str(key).strip(), str(value).strip(), now, now))

def memory_delete(category, key):
    with get_conn() as conn:
        return conn.execute("DELETE FROM memory WHERE category=? AND key=?", (category, key)).rowcount > 0

def memory_load_all():
    with get_conn() as conn:
        rows = conn.execute("SELECT category,key,value,updated FROM memory ORDER BY category,key").fetchall()
    if not rows: return "(no memories stored yet)"
    lines=[]; cur_cat=None
    for cat,key,val,updated in rows:
        if cat!=cur_cat: lines.append(f"\n[{cat.upper()}]"); cur_cat=cat
        lines.append(f"  {key}: {val}  (updated {updated[:10]})")
    return "\n".join(lines).strip()

def history_append(chat_id, role, content):
    now = datetime.utcnow().isoformat()
    with get_conn() as conn:
        conn.execute("INSERT INTO history(chat_id,role,content,ts) VALUES(?,?,?,?)", (chat_id,role,content,now))
        conn.execute("DELETE FROM history WHERE chat_id=? AND id NOT IN (SELECT id FROM history WHERE chat_id=? ORDER BY id DESC LIMIT ?)", (chat_id,chat_id,MAX_HISTORY))

def history_get(chat_id):
    with get_conn() as conn:
        rows = conn.execute("SELECT role,content FROM history WHERE chat_id=? ORDER BY id",(chat_id,)).fetchall()
    return [{"role":r,"content":c} for r,c in rows]

def get_google_creds(token_file=None):
    path = token_file or GOOGLE_TOKEN
    creds = Credentials.from_authorized_user_file(path, GOOGLE_SCOPES)
    if creds and creds.expired and creds.refresh_token:
        creds.refresh(Request())
        with open(path,'w') as f: f.write(creds.to_json())
    return creds

def gmail_get_unread(max_results=10, token_file=None):
    try:
        svc = build('gmail','v1',credentials=get_google_creds(token_file))
        msgs = svc.users().messages().list(userId='me',labelIds=['INBOX','UNREAD'],maxResults=max_results).execute().get('messages',[])
        if not msgs: return "No unread emails."
        out=[]
        for msg in msgs:
            m=svc.users().messages().get(userId='me',id=msg['id'],format='metadata',metadataHeaders=['From','Subject','Date']).execute()
            h={x['name']:x['value'] for x in m['payload']['headers']}
            out.append(f"From: {h.get('From','?')}\nSubject: {h.get('Subject','?')}\nDate: {h.get('Date','?')}\nPreview: {m.get('snippet','')[:100]}\nID: {msg['id']}")
        label="durginfamily@gmail.com" if token_file==FAMILY_TOKEN else "seandurgin@gmail.com"
        return f"Unread in {label} ({len(msgs)}):\n\n"+"\n---\n".join(out)
    except Exception as e: return f"Gmail error: {e}"

def gmail_read_message(message_id, token_file=None):
    try:
        svc=build('gmail','v1',credentials=get_google_creds(token_file))
        m=svc.users().messages().get(userId='me',id=message_id,format='full').execute()
        h={x['name']:x['value'] for x in m['payload']['headers']}
        body=""
        html_body=""
        # Extract text and html parts recursively
        def extract_parts(payload):
            nonlocal body, html_body
            if 'parts' in payload:
                for p in payload['parts']:
                    extract_parts(p)
            else:
                mime=payload.get('mimeType','')
                data=payload.get('body',{}).get('data','')
                if data:
                    decoded=base64.urlsafe_b64decode(data).decode('utf-8',errors='replace')
                    if mime=='text/plain' and not body:
                        body=decoded
                    elif mime=='text/html' and not html_body:
                        html_body=decoded
        extract_parts(m['payload'])
        # If no plain text, convert HTML
        if not body and html_body:
            import re
            # Extract view-in-browser links
            vib_links=re.findall(r"href=[\x22\x27]([^\x22\x27]*(?:view|browser|online)[^\x22\x27]*)[\x22\x27]",html_body,re.I)
            # Strip HTML tags
            body=re.sub(r'<style[^>]*>.*?</style>',' ',html_body,flags=re.DOTALL)
            body=re.sub(r'<script[^>]*>.*?</script>',' ',body,flags=re.DOTALL)
            body=re.sub(r'<[^>]+>',' ',body)
            body=re.sub(r'\s+',' ',body).strip()
            if vib_links:
                body=f"[View in browser: {vib_links[0]}]\n\n"+body
        result=f"From: {h.get('From','?')}\nSubject: {h.get('Subject','?')}\nDate: {h.get('Date','?')}\n\n{body[:2500]}"
        return result
    except Exception as e: return f"Error: {e}"

def gmail_send(to, subject, body, token_file=None):
    try:
        svc=build('gmail','v1',credentials=get_google_creds(token_file))
        msg=MIMEText(body); msg['to']=to; msg['subject']=subject
        svc.users().messages().send(userId='me',body={'raw':base64.urlsafe_b64encode(msg.as_bytes()).decode()}).execute()
        return f"Email sent to {to}."
    except Exception as e: return f"Failed: {e}"

def gmail_search(query, max_results=10, token_file=None):
    try:
        svc=build('gmail','v1',credentials=get_google_creds(token_file))
        msgs=svc.users().messages().list(userId='me',q=query,maxResults=max_results).execute().get('messages',[])
        if not msgs: return f"No emails found matching: {query}"
        out=[]
        for msg in msgs:
            m=svc.users().messages().get(userId='me',id=msg['id'],format='metadata',metadataHeaders=['From','Subject','Date']).execute()
            h={x['name']:x['value'] for x in m['payload']['headers']}
            out.append("From: "+h.get("From","?")+chr(10)+"Subject: "+h.get("Subject","?")+chr(10)+"Date: "+h.get("Date","?")+chr(10)+"Preview: "+m.get("snippet","")[:150]+chr(10)+"ID: "+msg["id"])
        label="durginfamily@gmail.com" if token_file==FAMILY_TOKEN else "seandurgin@gmail.com"
        return f"Search results in {label} for '{query}' ({len(msgs)}):\n\n" + "\n---\n".join(out)
    except Exception as e: return f"Gmail search error: {e}"

def calendar_get_upcoming(max_results=10):
    try:
        svc=build('calendar','v3',credentials=get_google_creds())
        events=svc.events().list(calendarId='primary',timeMin=datetime.utcnow().isoformat()+'Z',maxResults=max_results,singleEvents=True,orderBy='startTime').execute().get('items',[])
        if not events: return "No upcoming events."
        lines=[f"Upcoming ({len(events)}):"]
        for e in events: lines.append(f"- {e['start'].get('dateTime',e['start'].get('date','?'))}: {e.get('summary','No title')}")
        return "\n".join(lines)
    except Exception as e: return f"Calendar error: {e}"

def calendar_add_event(summary, start, end, description="", location=""):
    try:
        svc=build('calendar','v3',credentials=get_google_creds())
        event={'summary':summary,'start':{'dateTime':start,'timeZone':'America/New_York'},'end':{'dateTime':end,'timeZone':'America/New_York'}}
        if description: event['description']=description
        if location: event['location']=location
        c=svc.events().insert(calendarId='primary',body=event).execute()
        return f"Event created: {c.get('summary')} on {c['start'].get('dateTime','?')}"
    except Exception as e: return f"Failed: {e}"

def drive_search_files(query, max_results=5, token_file=None):
    try:
        svc=build('drive','v3',credentials=get_google_creds(token_file))
        files=svc.files().list(q=f"name contains '{query}' and trashed=false",pageSize=max_results,
            fields="files(id,name,mimeType,modifiedTime,size,webViewLink)").execute().get('files',[])
        if not files: return f"No files found: {query}"
        label="family Drive" if token_file==FAMILY_TOKEN else "Drive"
        lines=[f"{label} files matching '{query}':"]
        for f in files:
            kind="Folder" if "folder" in f['mimeType'] else "File"
            size=f"{int(f.get('size',0))//1024}KB" if f.get('size') else ""
            lines.append(f"- [{kind}] {f['name']} — {f.get('modifiedTime','?')[:10]} {size}")
            lines.append(f"  ID: {f['id']}")
            if f.get('webViewLink'): lines.append(f"  {f['webViewLink']}")
        return "\n".join(lines)
    except Exception as e: return f"Drive search error: {e}"

def drive_list(folder_id="root", max_results=20, token_file=None):
    try:
        svc=build('drive','v3',credentials=get_google_creds(token_file))
        files=svc.files().list(q=f"'{folder_id}' in parents and trashed=false",pageSize=max_results,
            fields="files(id,name,mimeType,modifiedTime,size,webViewLink)").execute().get('files',[])
        if not files: return "No files found in folder."
        label="family Drive" if token_file==FAMILY_TOKEN else "Drive"
        lines=[f"{label} folder contents ({len(files)}):"]
        for f in files:
            kind="Folder" if "folder" in f['mimeType'] else "File"
            size=f"{int(f.get('size',0))//1024}KB" if f.get('size') else ""
            lines.append(f"- [{kind}] {f['name']} — {f.get('modifiedTime','?')[:10]} {size}")
            lines.append(f"  ID: {f['id']}")
        return "\n".join(lines)
    except Exception as e: return f"Drive list error: {e}"

def drive_read_file(file_id, token_file=None):
    try:
        svc=build('drive','v3',credentials=get_google_creds(token_file))
        meta=svc.files().get(fileId=file_id,fields="name,mimeType").execute()
        filename=meta['name']; mime=meta['mimeType']
        if 'google-apps' in mime:
            export_map={'application/vnd.google-apps.document':'text/plain','application/vnd.google-apps.spreadsheet':'text/csv','application/vnd.google-apps.presentation':'text/plain'}
            export_mime=export_map.get(mime,'text/plain')
            data=svc.files().export(fileId=file_id,mimeType=export_mime).execute()
            return data.decode('utf-8',errors='replace')[:4000]
        else:
            from googleapiclient.http import MediaIoBaseDownload
            buf=io.BytesIO()
            dl=MediaIoBaseDownload(buf,svc.files().get_media(fileId=file_id))
            done=False
            while not done: _,done=dl.next_chunk()
            return extract_file_text(buf.getvalue(),filename)
    except Exception as e: return f"Drive read error: {e}"

def drive_move_file(file_id, new_folder_id, token_file=None):
    try:
        svc=build('drive','v3',credentials=get_google_creds(token_file))
        f=svc.files().get(fileId=file_id,fields="parents").execute()
        prev=",".join(f.get('parents',[]))
        svc.files().update(fileId=file_id,addParents=new_folder_id,removeParents=prev,fields="id,parents").execute()
        return f"File moved."
    except Exception as e: return f"Drive move error: {e}"

def drive_rename_file(file_id, new_name, token_file=None):
    try:
        svc=build('drive','v3',credentials=get_google_creds(token_file))
        svc.files().update(fileId=file_id,body={"name":new_name}).execute()
        return f"Renamed to '{new_name}'."
    except Exception as e: return f"Drive rename error: {e}"

def drive_delete_file(file_id, token_file=None):
    try:
        svc=build('drive','v3',credentials=get_google_creds(token_file))
        svc.files().delete(fileId=file_id).execute()
        return "File deleted."
    except Exception as e: return f"Drive delete error: {e}"

def drive_create_folder(name, parent_id="root", token_file=None):
    try:
        svc=build('drive','v3',credentials=get_google_creds(token_file))
        meta={"name":name,"mimeType":"application/vnd.google-apps.folder"}
        if parent_id!="root": meta["parents"]=[parent_id]
        f=svc.files().create(body=meta,fields="id,name,webViewLink").execute()
        return f"Folder '{name}' created. ID: {f['id']}\nLink: {f.get('webViewLink','')}"
    except Exception as e: return f"Drive create folder error: {e}"

def drive_upload_file(filename, content_text, folder_id="root", token_file=None):
    try:
        from googleapiclient.http import MediaIoBaseUpload
        svc=build('drive','v3',credentials=get_google_creds(token_file))
        meta={"name":filename}
        if folder_id!="root": meta["parents"]=[folder_id]
        media=MediaIoBaseUpload(io.BytesIO(content_text.encode('utf-8')),mimetype='text/plain')
        f=svc.files().create(body=meta,media_body=media,fields="id,name,webViewLink").execute()
        return f"Uploaded '{filename}'.\nID: {f['id']}\nLink: {f.get('webViewLink','')}"
    except Exception as e: return f"Drive upload error: {e}"

def contacts_search(query, max_results=5):
    try:
        svc=build('people','v1',credentials=get_google_creds())
        results=svc.people().searchContacts(query=query,readMask='names,emailAddresses,phoneNumbers,organizations',pageSize=max_results).execute().get('results',[])
        if not results: return f"No contacts: {query}"
        lines=[f"Contacts matching '{query}':"]
        for p in results:
            person=p.get('person',{})
            name=person.get('names',[{}])[0].get('displayName','Unknown')
            emails=[e['value'] for e in person.get('emailAddresses',[])]
            phones=[ph['value'] for ph in person.get('phoneNumbers',[])]
            lines.append(f"\n{name}")
            if emails: lines.append(f"  {', '.join(emails)}")
            if phones: lines.append(f"  {', '.join(phones)}")
        return "\n".join(lines)
    except Exception as e: return f"Contacts error: {e}"

def tasks_list_tasklists():
    try:
        svc=build('tasks','v1',credentials=get_google_creds())
        result=svc.tasklists().list().execute()
        lists=result.get('items',[])
        if not lists: return "No task lists found."
        lines=["Your task lists:"]
        for tl in lists: lines.append(f"- {tl['title']} (ID: {tl['id']})")
        return "\n".join(lines)
    except Exception as e: return f"Tasks error: {e}"

def tasks_list(tasklist_id="@default", show_completed=False):
    try:
        svc=build('tasks','v1',credentials=get_google_creds())
        result=svc.tasks().list(tasklist=tasklist_id,showCompleted=show_completed,showHidden=False).execute()
        tasks=result.get('items',[])
        if not tasks: return "No tasks found."
        lines=[f"Tasks ({len(tasks)}):"]
        for t in tasks:
            status="✅" if t.get('status')=='completed' else "⬜"
            due=f" - due {t['due'][:10]}" if t.get('due') else ""
            notes=f"\n  {t['notes']}" if t.get('notes') else ""
            lines.append(f"{status} {t['title']}{due} (ID: {t['id']}){notes}")
        return "\n".join(lines)
    except Exception as e: return f"Tasks error: {e}"

def tasks_add(title, notes="", due=None, tasklist_id="@default"):
    try:
        svc=build('tasks','v1',credentials=get_google_creds())
        body={"title":title}
        if notes: body["notes"]=notes
        if due: body["due"]=due+"T00:00:00.000Z" if "T" not in due else due
        task=svc.tasks().insert(tasklist=tasklist_id,body=body).execute()
        return f"Task added: '{task['title']}' (ID: {task['id']})"
    except Exception as e: return f"Tasks error: {e}"

def tasks_complete(task_id, tasklist_id="@default"):
    try:
        svc=build('tasks','v1',credentials=get_google_creds())
        task=svc.tasks().get(tasklist=tasklist_id,task=task_id).execute()
        task['status']='completed'
        svc.tasks().update(tasklist=tasklist_id,task=task_id,body=task).execute()
        return f"Task completed: '{task['title']}'"
    except Exception as e: return f"Tasks error: {e}"

def tasks_delete(task_id, tasklist_id="@default"):
    try:
        svc=build('tasks','v1',credentials=get_google_creds())
        svc.tasks().delete(tasklist=tasklist_id,task=task_id).execute()
        return "Task deleted."
    except Exception as e: return f"Tasks error: {e}"

def ms_get_token():
    with open(MS_TOKEN) as f: td=json.load(f)
    app=msal.PublicClientApplication(MS_CLIENT_ID,authority=MS_AUTHORITY)
    result=None
    if 'refresh_token' in td: result=app.acquire_token_by_refresh_token(td['refresh_token'],scopes=MS_SCOPES)
    if not result or 'access_token' not in result: raise Exception("Could not refresh Microsoft token.")
    td.update(result)
    with open(MS_TOKEN,'w') as f: json.dump(td,f)
    return result['access_token']

def ms_get(path, params=None):
    r=requests.get(f"{GRAPH_BASE}{path}",headers={"Authorization":f"Bearer {ms_get_token()}"},params=params,timeout=15)
    r.raise_for_status(); return r.json()

def onenote_list_notebooks():
    try:
        nbs=ms_get("/me/onenote/notebooks").get('value',[])
        if not nbs: return "No OneNote notebooks found."
        return "Your OneNote notebooks:\n"+"\n".join(f"- {nb['displayName']} (ID: {nb['id']})" for nb in nbs)
    except Exception as e: return f"OneNote error: {e}"

def onenote_list_sections(notebook_name=None):
    try:
        if notebook_name:
            nbs=ms_get("/me/onenote/notebooks").get('value',[])
            nb=next((n for n in nbs if notebook_name.lower() in n['displayName'].lower()),None)
            if not nb: return f"Notebook not found: {notebook_name}"
            sections=ms_get(f"/me/onenote/notebooks/{nb['id']}/sections").get('value',[])
        else:
            sections=ms_get("/me/onenote/sections").get('value',[])
        if not sections: return "No sections found."
        return "Sections:\n"+"\n".join(f"- {s['displayName']} (ID: {s['id']})" for s in sections)
    except Exception as e: return f"OneNote error: {e}"

def onenote_list_pages(section_id, max_results=20):
    try:
        pages=ms_get(f"/me/onenote/sections/{section_id}/pages",params={"$top":max_results,"$orderby":"lastModifiedDateTime desc","$select":"title,lastModifiedDateTime,id"}).get('value',[])
        if not pages: return "No pages found."
        lines=[f"Pages ({len(pages)}):"]
        for p in pages: lines.append(f"- {p['title']} — {p.get('lastModifiedDateTime','?')[:10]} (ID: {p['id']})")
        return "\n".join(lines)
    except Exception as e: return f"OneNote error: {e}"

def onenote_recent_pages(max_results=10):
    try:
        pages=ms_get("/me/onenote/pages",params={"$top":max_results,"$orderby":"lastModifiedDateTime desc","$select":"title,lastModifiedDateTime,parentSection,id"}).get('value',[])
        if not pages: return "No recent pages."
        lines=[f"Recent pages ({len(pages)}):"]
        for p in pages: lines.append(f"- {p['title']} [{p.get('parentSection',{}).get('displayName','?')}] - {p.get('lastModifiedDateTime','?')[:10]} (ID: {p['id']})")
        return "\n".join(lines)
    except Exception as e: return f"OneNote error: {e}"

def onenote_search_pages(query, max_results=5):
    try:
        pages=ms_get("/me/onenote/pages",params={"$top":max_results,"$search":query,"$select":"title,lastModifiedDateTime,parentSection,id"}).get('value',[])
        if not pages: return f"No pages matching: {query}"
        lines=[f"Pages matching '{query}':"]
        for p in pages: lines.append(f"- {p['title']} [{p.get('parentSection',{}).get('displayName','?')}] - {p.get('lastModifiedDateTime','?')[:10]} (ID: {p['id']})")
        return "\n".join(lines)
    except Exception as e: return f"OneNote search error: {e}"

def onenote_get_page(page_id):
    try:
        r=requests.get(f"{GRAPH_BASE}/me/onenote/pages/{page_id}/content",headers={"Authorization":f"Bearer {ms_get_token()}"},timeout=15)
        r.raise_for_status()
        text=re.sub(r'\s+',' ',re.sub(r'<[^>]+>',' ',r.text)).strip()
        return text[:2000]+"\n[truncated]" if len(text)>2000 else text
    except Exception as e: return f"Error reading page: {e}"

def onenote_create_page(section_id, title, content):
    try:
        html=f"<!DOCTYPE html><html><head><title>{title}</title></head><body><h1>{title}</h1><p>{content}</p></body></html>"
        r=requests.post(f"{GRAPH_BASE}/me/onenote/sections/{section_id}/pages",headers={"Authorization":f"Bearer {ms_get_token()}","Content-Type":"application/xhtml+xml"},data=html.encode('utf-8'),timeout=15)
        r.raise_for_status(); return f"Page created: {title}"
    except Exception as e: return f"Failed: {e}"

def onedrive_search(query, max_results=10):
    try:
        token=ms_get_token()
        r=requests.get(f"{GRAPH_BASE}/me/drive/search(q='{query}')",headers={"Authorization":f"Bearer {token}"},
            params={"$top":max_results,"$select":"name,size,lastModifiedDateTime,webUrl,file,folder,id"},timeout=15)
        r.raise_for_status()
        items=r.json().get('value',[])
        if not items: return f"No OneDrive files found: {query}"
        lines=[f"OneDrive files matching '{query}' ({len(items)}):"]
        for item in items:
            kind="Folder" if "folder" in item else "File"
            size=f"{item.get('size',0)//1024}KB" if "file" in item else ""
            lines.append(f"- [{kind}] {item['name']} — {item.get('lastModifiedDateTime','?')[:10]} {size}")
            lines.append(f"  ID: {item['id']}")
            if item.get('webUrl'): lines.append(f"  {item['webUrl']}")
        return "\n".join(lines)
    except Exception as e: return f"OneDrive search error: {e}"

def onedrive_list(folder_id="root", max_results=20):
    try:
        token=ms_get_token()
        if folder_id in ("root","/"): url=f"{GRAPH_BASE}/me/drive/root/children"
        else: url=f"{GRAPH_BASE}/me/drive/items/{folder_id}/children"
        r=requests.get(url,headers={"Authorization":f"Bearer {token}"},
            params={"$top":max_results,"$select":"name,size,lastModifiedDateTime,webUrl,file,folder,id"},timeout=15)
        r.raise_for_status()
        items=r.json().get('value',[])
        if not items: return "No files found in OneDrive folder."
        lines=[f"OneDrive folder contents ({len(items)}):"]
        for item in items:
            kind="Folder" if "folder" in item else "File"
            size=f"{item.get('size',0)//1024}KB" if "file" in item else ""
            lines.append(f"- [{kind}] {item['name']} — {item.get('lastModifiedDateTime','?')[:10]} {size}")
            lines.append(f"  ID: {item['id']}")
        return "\n".join(lines)
    except Exception as e: return f"OneDrive list error: {e}"

def onedrive_read(item_id):
    try:
        token=ms_get_token()
        meta=requests.get(f"{GRAPH_BASE}/me/drive/items/{item_id}",headers={"Authorization":f"Bearer {token}"},timeout=15)
        meta.raise_for_status(); filename=meta.json().get('name','unknown')
        r=requests.get(f"{GRAPH_BASE}/me/drive/items/{item_id}/content",headers={"Authorization":f"Bearer {token}"},timeout=30,allow_redirects=True)
        r.raise_for_status()
        return extract_file_text(r.content,filename)
    except Exception as e: return f"OneDrive read error: {e}"

def onedrive_upload(filename, content_text, folder_path="/"):
    try:
        token=ms_get_token()
        if folder_path in ("/","root"): url=f"{GRAPH_BASE}/me/drive/root:/{filename}:/content"
        else: url=f"{GRAPH_BASE}/me/drive/root:/{folder_path}/{filename}:/content"
        r=requests.put(url,headers={"Authorization":f"Bearer {token}","Content-Type":"text/plain"},data=content_text.encode('utf-8'),timeout=30)
        r.raise_for_status(); result=r.json()
        return f"Uploaded '{filename}' to OneDrive.\nLink: {result.get('webUrl','')}"
    except Exception as e: return f"OneDrive upload error: {e}"

def onedrive_delete(item_id):
    try:
        token=ms_get_token()
        r=requests.delete(f"{GRAPH_BASE}/me/drive/items/{item_id}",headers={"Authorization":f"Bearer {token}"},timeout=15)
        if r.status_code==204: return "File deleted from OneDrive."
        r.raise_for_status(); return "Deleted."
    except Exception as e: return f"OneDrive delete error: {e}"

def outlook_unread(max_results=10):
    try:
        msgs=ms_get("/me/mailFolders/inbox/messages",params={"$filter":"isRead eq false","$top":max_results,"$select":"subject,from,receivedDateTime,bodyPreview,id","$orderby":"receivedDateTime desc"}).get("value",[])
        if not msgs: return "No unread emails in Outlook."
        lines=[f"Unread in seandurgin@live.com ({len(msgs)}):"]
        for m in msgs:
            sender=m.get("from",{}).get("emailAddress",{})
            lines.append(f"From: {sender.get('name','?')} <{sender.get('address','?')}>")
            lines.append(f"Subject: {m.get('subject','?')}")
            lines.append(f"Date: {m.get('receivedDateTime','?')[:10]}")
            lines.append(f"Preview: {m.get('bodyPreview','')[:150]}")
            lines.append(f"ID: {m.get('id','?')}")
            lines.append("---")
        return "\n".join(lines)
    except Exception as e: return f"Outlook error: {e}"

def outlook_read(message_id):
    try:
        m=ms_get(f"/me/messages/{message_id}",params={"$select":"subject,from,receivedDateTime,body"})
        sender=m.get("from",{}).get("emailAddress",{})
        body=m.get("body",{})
        content_type=body.get("contentType","text")
        text=body.get("content","")
        if content_type=="html":
            import re
            text=re.sub(r"<style[^>]*>.*?</style>"," ",text,flags=re.DOTALL)
            text=re.sub(r"<[^>]+>"," ",text)
            text=re.sub(r"\s+"," ",text).strip()
        return f"From: {sender.get('name','?')} <{sender.get('address','?')}>\nSubject: {m.get('subject','?')}\nDate: {m.get('receivedDateTime','?')[:10]}\n\n{text[:2500]}"
    except Exception as e: return f"Outlook read error: {e}"

def outlook_search(query, max_results=10):
    try:
        msgs=ms_get("/me/messages",params={"$search":f'"{query}"',"$top":max_results,"$select":"subject,from,receivedDateTime,bodyPreview,id"}).get("value",[])
        if not msgs: return f"No Outlook emails found matching: {query}"
        lines=[f"Outlook search results for '{query}' ({len(msgs)}):"]
        for m in msgs:
            sender=m.get("from",{}).get("emailAddress",{})
            lines.append(f"From: {sender.get('name','?')} <{sender.get('address','?')}>")
            lines.append(f"Subject: {m.get('subject','?')}")
            lines.append(f"Date: {m.get('receivedDateTime','?')[:10]}")
            lines.append(f"Preview: {m.get('bodyPreview','')[:150]}")
            lines.append(f"ID: {m.get('id','?')}")
            lines.append("---")
        return "\n".join(lines)
    except Exception as e: return f"Outlook search error: {e}"

def outlook_send(to, subject, body):
    try:
        import requests as req
        token=ms_get_token()
        msg={"message":{"subject":subject,"body":{"contentType":"Text","content":body},"toRecipients":[{"emailAddress":{"address":to}}]}}
        r=req.post(f"{GRAPH_BASE}/me/sendMail",headers={"Authorization":f"Bearer {token}","Content-Type":"application/json"},json=msg,timeout=15)
        r.raise_for_status()
        return f"Email sent to {to} from seandurgin@live.com."
    except Exception as e: return f"Outlook send error: {e}"

async def brave_search(query, count=5):
    if not BRAVE_KEY: return "Web search not configured."
    try:
        async with httpx.AsyncClient(timeout=10) as client:
            r=await client.get("https://api.search.brave.com/res/v1/web/search",headers={"Accept":"application/json","X-Subscription-Token":BRAVE_KEY},params={"q":query,"count":count,"text_decorations":False})
            r.raise_for_status(); results=r.json().get("web",{}).get("results",[])
        if not results: return f"No results: {query}"
        lines=[f"Results for: {query}\n"]
        for i,res in enumerate(results[:count],1): lines.append(f"{i}. {res.get('title','')}\n   {res.get('url','')}\n   {res.get('description','')}\n")
        return "\n".join(lines)
    except Exception as e: return f"Search failed: {e}"

async def fetch_url(url):
    try:
        async with httpx.AsyncClient(timeout=15,follow_redirects=True) as client:
            r=await client.get(url,headers={"User-Agent":"Mozilla/5.0 (compatible; Clawdia/1.0)"})
            r.raise_for_status()
            text=re.sub(r'<script[^>]*>.*?</script>',' ',r.text,flags=re.DOTALL)
            text=re.sub(r'<style[^>]*>.*?</style>',' ',text,flags=re.DOTALL)
            text=re.sub(r'<[^>]+>',' ',text)
            text=re.sub(r'\s+',' ',text).strip()
            if len(text)>2500: return text[:2500]+"\n[truncated]"
            return f"Content from {url}:\n\n{text}"
    except Exception as e: return f"Failed to fetch {url}: {e}"

def extract_file_text(file_bytes, filename):
    ext=filename.lower().rsplit('.',1)[-1]
    # Size check — skip files over 8MB
    if len(file_bytes) > 8 * 1024 * 1024:
        return f"File too large ({len(file_bytes)//1024//1024}MB) — skipping. Max 8MB."
    try:
        if ext=='pdf':
            import PyPDF2
            reader=PyPDF2.PdfReader(io.BytesIO(file_bytes))
            text=""
            for page in reader.pages: text+=page.extract_text() or ""
            text=text.strip()
            if not text:
                try:
                    import pytesseract
                    from pdf2image import convert_from_bytes
                    images=convert_from_bytes(file_bytes,dpi=200)
                    for img in images[:5]: text+=pytesseract.image_to_string(img)
                    text=text.strip()
                except Exception as ocr_err:
                    return f"PDF has no text layer and OCR failed: {ocr_err}"
            return text[:4000] or "Could not extract text from PDF."
        elif ext=='docx':
            try:
                import docx as docxlib
                d=docxlib.Document(io.BytesIO(file_bytes))
                parts=[]
                for p in d.paragraphs:
                    if p.text.strip(): parts.append(p.text.strip())
                for table in d.tables:
                    for row in table.rows:
                        row_text="\t".join(c.text.strip() for c in row.cells if c.text.strip())
                        if row_text: parts.append(row_text)
                return "\n".join(parts)[:4000] or "Empty document."
            except Exception as e:
                return f"DOCX read error: {e}"
        elif ext in ('xlsx','xls'):
            import openpyxl
            wb=openpyxl.load_workbook(io.BytesIO(file_bytes),read_only=True,data_only=True)
            lines=[]
            for sheet in wb.sheetnames[:3]:
                ws=wb[sheet]; lines.append(f"[Sheet: {sheet}]")
                for i,row in enumerate(ws.iter_rows(values_only=True)):
                    if any(c is not None for c in row):
                        lines.append("\t".join(str(c) if c is not None else "" for c in row))
                    if i>=50: lines.append("...(truncated)"); break
            return "\n".join(lines)[:4000] or "Empty spreadsheet."
        elif ext=='csv':
            import csv
            raw=file_bytes.decode('utf-8',errors='replace')
            lines=[]
            for i,row in enumerate(csv.reader(io.StringIO(raw))):
                lines.append("\t".join(row))
                if i>=50: lines.append("...(truncated)"); break
            return "\n".join(lines)[:4000]
        elif ext in ('txt','md','json','py','js','html','css'):
            return file_bytes.decode('utf-8',errors='replace')[:4000]
        elif ext in ('jpg','jpeg','png','gif','webp','bmp'):
            return f"IMAGE_FILE:{base64.b64encode(file_bytes).decode('utf-8')[:100000]}"
        else:
            return f"Unsupported file type: .{ext}. I can read PDF, DOCX, XLSX, CSV, TXT, images."
    except Exception as e: return f"Error reading {filename}: {e}"

TOOLS = [
    {"name":"save_memory","description":"Save or update a fact about Sean. Categories: personal, health, preferences, work, family, notes.","input_schema":{"type":"object","properties":{"category":{"type":"string"},"key":{"type":"string"},"value":{"type":"string"}},"required":["category","key","value"]}},
    {"name":"delete_memory","description":"Delete a memory entry.","input_schema":{"type":"object","properties":{"category":{"type":"string"},"key":{"type":"string"}},"required":["category","key"]}},
    {"name":"web_search","description":"Search the web for current information.","input_schema":{"type":"object","properties":{"query":{"type":"string"},"count":{"type":"integer","default":5}},"required":["query"]}},
    {"name":"fetch_url","description":"Fetch and read the content of any URL.","input_schema":{"type":"object","properties":{"url":{"type":"string"}},"required":["url"]}},
    {"name":"gmail_unread","description":"Get unread emails from seandurgin@gmail.com.","input_schema":{"type":"object","properties":{"max_results":{"type":"integer","default":10}}}},
    {"name":"gmail_read","description":"Read a specific email from seandurgin@gmail.com by ID.","input_schema":{"type":"object","properties":{"message_id":{"type":"string"}},"required":["message_id"]}},
    {"name":"gmail_send","description":"Send email from seandurgin@gmail.com. ALWAYS confirm with Sean first.","input_schema":{"type":"object","properties":{"to":{"type":"string"},"subject":{"type":"string"},"body":{"type":"string"}},"required":["to","subject","body"]}},
    {"name":"gmail_search","description":"Search all emails in seandurgin@gmail.com using Gmail query syntax e.g. from:usps.com, subject:invoice, after:2026/01/01.","input_schema":{"type":"object","properties":{"query":{"type":"string"},"max_results":{"type":"integer","default":10}},"required":["query"]}},
    {"name":"family_gmail_search","description":"Search all emails in durginfamily@gmail.com using Gmail query syntax.","input_schema":{"type":"object","properties":{"query":{"type":"string"},"max_results":{"type":"integer","default":10}},"required":["query"]}},
    {"name":"family_gmail_unread","description":"Get unread emails from durginfamily@gmail.com.","input_schema":{"type":"object","properties":{"max_results":{"type":"integer","default":10}}}},
    {"name":"family_gmail_read","description":"Read a specific email from durginfamily@gmail.com by ID.","input_schema":{"type":"object","properties":{"message_id":{"type":"string"}},"required":["message_id"]}},
    {"name":"family_gmail_send","description":"Send email from durginfamily@gmail.com. ALWAYS confirm with Sean first.","input_schema":{"type":"object","properties":{"to":{"type":"string"},"subject":{"type":"string"},"body":{"type":"string"}},"required":["to","subject","body"]}},
    {"name":"calendar_upcoming","description":"Get Sean's upcoming Google Calendar events.","input_schema":{"type":"object","properties":{"max_results":{"type":"integer","default":10}}}},
    {"name":"calendar_add","description":"Add event to Google Calendar. ISO 8601 format for start/end.","input_schema":{"type":"object","properties":{"summary":{"type":"string"},"start":{"type":"string"},"end":{"type":"string"},"description":{"type":"string"},"location":{"type":"string"}},"required":["summary","start","end"]}},
    {"name":"drive_search","description":"Search files in seandurgin Google Drive by name.","input_schema":{"type":"object","properties":{"query":{"type":"string"},"max_results":{"type":"integer","default":5}},"required":["query"]}},
    {"name":"drive_list","description":"List contents of a seandurgin Google Drive folder. Use 'root' for root.","input_schema":{"type":"object","properties":{"folder_id":{"type":"string","default":"root"},"max_results":{"type":"integer","default":20}}}},
    {"name":"drive_read","description":"Read a file from seandurgin Google Drive by ID.","input_schema":{"type":"object","properties":{"file_id":{"type":"string"}},"required":["file_id"]}},
    {"name":"drive_move","description":"Move a file in seandurgin Google Drive. Confirm with Sean first.","input_schema":{"type":"object","properties":{"file_id":{"type":"string"},"new_folder_id":{"type":"string"}},"required":["file_id","new_folder_id"]}},
    {"name":"drive_rename","description":"Rename a file in seandurgin Google Drive.","input_schema":{"type":"object","properties":{"file_id":{"type":"string"},"new_name":{"type":"string"}},"required":["file_id","new_name"]}},
    {"name":"drive_delete","description":"Delete a file from seandurgin Google Drive. ALWAYS confirm with Sean first.","input_schema":{"type":"object","properties":{"file_id":{"type":"string"}},"required":["file_id"]}},
    {"name":"drive_create_folder","description":"Create a folder in seandurgin Google Drive.","input_schema":{"type":"object","properties":{"name":{"type":"string"},"parent_id":{"type":"string","default":"root"}},"required":["name"]}},
    {"name":"drive_upload","description":"Upload a text file to seandurgin Google Drive. Confirm with Sean first.","input_schema":{"type":"object","properties":{"filename":{"type":"string"},"content_text":{"type":"string"},"folder_id":{"type":"string","default":"root"}},"required":["filename","content_text"]}},
    {"name":"family_drive_search","description":"Search files in durginfamily Google Drive by name.","input_schema":{"type":"object","properties":{"query":{"type":"string"},"max_results":{"type":"integer","default":5}},"required":["query"]}},
    {"name":"family_drive_list","description":"List contents of a durginfamily Google Drive folder.","input_schema":{"type":"object","properties":{"folder_id":{"type":"string","default":"root"},"max_results":{"type":"integer","default":20}}}},
    {"name":"family_drive_read","description":"Read a file from durginfamily Google Drive by ID.","input_schema":{"type":"object","properties":{"file_id":{"type":"string"}},"required":["file_id"]}},
    {"name":"contacts_search","description":"Search Sean's Google Contacts.","input_schema":{"type":"object","properties":{"query":{"type":"string"},"max_results":{"type":"integer","default":5}},"required":["query"]}},
    {"name":"tasks_list_tasklists","description":"List all of Sean's Google Task lists.","input_schema":{"type":"object","properties":{}}},
    {"name":"tasks_list","description":"List tasks in a Google Task list.","input_schema":{"type":"object","properties":{"tasklist_id":{"type":"string","default":"@default"},"show_completed":{"type":"boolean","default":False}}}},
    {"name":"tasks_add","description":"Add a task to Google Tasks.","input_schema":{"type":"object","properties":{"title":{"type":"string"},"notes":{"type":"string"},"due":{"type":"string"},"tasklist_id":{"type":"string","default":"@default"}},"required":["title"]}},
    {"name":"tasks_complete","description":"Mark a Google Task as complete.","input_schema":{"type":"object","properties":{"task_id":{"type":"string"},"tasklist_id":{"type":"string","default":"@default"}},"required":["task_id"]}},
    {"name":"tasks_delete","description":"Delete a Google Task.","input_schema":{"type":"object","properties":{"task_id":{"type":"string"},"tasklist_id":{"type":"string","default":"@default"}},"required":["task_id"]}},
    {"name":"outlook_unread","description":"Get unread emails from seandurgin@live.com (personal Microsoft/Outlook account).","input_schema":{"type":"object","properties":{"max_results":{"type":"integer","default":10}}}},
    {"name":"outlook_read","description":"Read a specific email from seandurgin@live.com by message ID.","input_schema":{"type":"object","properties":{"message_id":{"type":"string"}},"required":["message_id"]}},
    {"name":"outlook_search","description":"Search emails in seandurgin@live.com.","input_schema":{"type":"object","properties":{"query":{"type":"string"},"max_results":{"type":"integer","default":10}},"required":["query"]}},
    {"name":"outlook_send","description":"Send email from seandurgin@live.com. ALWAYS confirm with Sean first.","input_schema":{"type":"object","properties":{"to":{"type":"string"},"subject":{"type":"string"},"body":{"type":"string"}},"required":["to","subject","body"]}},
    {"name":"onenote_notebooks","description":"List all of Sean's OneNote notebooks.","input_schema":{"type":"object","properties":{}}},
    {"name":"onenote_sections","description":"List sections in a OneNote notebook.","input_schema":{"type":"object","properties":{"notebook_name":{"type":"string"}}}},
    {"name":"onenote_list_pages","description":"List pages in a OneNote section by section ID.","input_schema":{"type":"object","properties":{"section_id":{"type":"string"},"max_results":{"type":"integer","default":20}},"required":["section_id"]}},
    {"name":"onenote_recent","description":"Get recently modified OneNote pages.","input_schema":{"type":"object","properties":{"max_results":{"type":"integer","default":10}}}},
    {"name":"onenote_search","description":"Search OneNote pages by keyword.","input_schema":{"type":"object","properties":{"query":{"type":"string"},"max_results":{"type":"integer","default":5}},"required":["query"]}},
    {"name":"onenote_read","description":"Read a OneNote page by ID.","input_schema":{"type":"object","properties":{"page_id":{"type":"string"}},"required":["page_id"]}},
    {"name":"onenote_create","description":"Create a new OneNote page in a section.","input_schema":{"type":"object","properties":{"section_id":{"type":"string"},"title":{"type":"string"},"content":{"type":"string"}},"required":["section_id","title","content"]}},
    {"name":"onedrive_search","description":"Search files in Sean's OneDrive.","input_schema":{"type":"object","properties":{"query":{"type":"string"},"max_results":{"type":"integer","default":10}},"required":["query"]}},
    {"name":"onedrive_list","description":"List contents of a OneDrive folder.","input_schema":{"type":"object","properties":{"folder_id":{"type":"string","default":"root"},"max_results":{"type":"integer","default":20}}}},
    {"name":"onedrive_read","description":"Read a OneDrive file by ID.","input_schema":{"type":"object","properties":{"item_id":{"type":"string"}},"required":["item_id"]}},
    {"name":"onedrive_upload","description":"Upload a file to OneDrive. Confirm with Sean first.","input_schema":{"type":"object","properties":{"filename":{"type":"string"},"content_text":{"type":"string"},"folder_path":{"type":"string","default":"/"}},"required":["filename","content_text"]}},
    {"name":"onedrive_delete","description":"Delete a OneDrive file. ALWAYS confirm with Sean first.","input_schema":{"type":"object","properties":{"item_id":{"type":"string"}},"required":["item_id"]}},
    {"name":"plaid_accounts","description":"Get current bank account balances for all connected accounts. Only call when Sean explicitly asks.","input_schema":{"type":"object","properties":{}}},
    {"name":"plaid_transactions","description":"Get recent transactions. Only call when Sean explicitly asks.","input_schema":{"type":"object","properties":{"days":{"type":"integer","default":30},"max_results":{"type":"integer","default":50}}}},
    {"name":"plaid_spending_by_category","description":"Summarize spending by category. Only call when Sean explicitly asks.","input_schema":{"type":"object","properties":{"days":{"type":"integer","default":30}}}},
    {"name":"plaid_debt_snapshot","description":"Get current debt balances and save a snapshot to memory for tracking over time. Only call when Sean explicitly asks.","input_schema":{"type":"object","properties":{}}},
    {"name":"reminder_set","description":"Set a reminder for Sean. remind_at must be ISO 8601 UTC datetime string e.g. 2026-04-15T14:00:00. Convert from Eastern time (UTC-4 in summer, UTC-5 in winter).","input_schema":{"type":"object","properties":{"message":{"type":"string","description":"What to remind Sean about"},"remind_at":{"type":"string","description":"ISO 8601 UTC datetime e.g. 2026-04-15T18:00:00"}},"required":["message","remind_at"]}},
    {"name":"reminder_list","description":"List all of Sean's pending reminders.","input_schema":{"type":"object","properties":{}}},
    {"name":"reminder_delete","description":"Delete a reminder by ID.","input_schema":{"type":"object","properties":{"reminder_id":{"type":"integer"}},"required":["reminder_id"]}},
]

async def describe_image_bytes(file_bytes, filename, question=""):
    """Pass image bytes to Claude vision and return description."""
    try:
        client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_KEY)
        img_b64 = base64.b64encode(file_bytes).decode('utf-8')
        ext = filename.lower().rsplit('.',1)[-1]
        media_map = {'jpg':'image/jpeg','jpeg':'image/jpeg','png':'image/png','gif':'image/gif','webp':'image/webp'}
        media_type = media_map.get(ext, 'image/jpeg')
        q = question if question else f"Describe this image in detail. If it contains text, read all of it. File: {filename}"
        response = await client.messages.create(
            model=MODEL, max_tokens=1024,
            messages=[{"role":"user","content":[
                {"type":"image","source":{"type":"base64","media_type":media_type,"data":img_b64}},
                {"type":"text","text":q}
            ]}]
        )
        return response.content[0].text
    except Exception as e:
        return f"Image vision error: {e}"

async def run_tool(name, inputs):
    if name=="save_memory":
        cat=str(inputs.get("category","")).strip()
        key=str(inputs.get("key","")).strip()
        val=str(inputs.get("value","")).strip()
        if not cat or not key or not val: return "Skipped: missing field"
        memory_save(cat,key,val); return f"Remembered: [{cat}] {key} = {val}"
    elif name=="delete_memory": return "Deleted." if memory_delete(inputs["category"],inputs["key"]) else "Not found."
    elif name=="web_search": return await brave_search(inputs["query"],inputs.get("count",5))
    elif name=="fetch_url": return await fetch_url(inputs["url"])
    elif name=="gmail_unread": return await asyncio.to_thread(gmail_get_unread,inputs.get("max_results",10))
    elif name=="gmail_read": return await asyncio.to_thread(gmail_read_message,inputs["message_id"])
    elif name=="gmail_send": return await asyncio.to_thread(gmail_send,inputs["to"],inputs["subject"],inputs["body"])
    elif name=="gmail_search": return await asyncio.to_thread(gmail_search,inputs["query"],inputs.get("max_results",10))
    elif name=="family_gmail_search": return await asyncio.to_thread(gmail_search,inputs["query"],inputs.get("max_results",10),FAMILY_TOKEN)
    elif name=="family_gmail_unread": return await asyncio.to_thread(gmail_get_unread,inputs.get("max_results",10),FAMILY_TOKEN)
    elif name=="family_gmail_read": return await asyncio.to_thread(gmail_read_message,inputs["message_id"],FAMILY_TOKEN)
    elif name=="family_gmail_send": return await asyncio.to_thread(gmail_send,inputs["to"],inputs["subject"],inputs["body"],FAMILY_TOKEN)
    elif name=="calendar_upcoming": return await asyncio.to_thread(calendar_get_upcoming,inputs.get("max_results",10))
    elif name=="calendar_add": return await asyncio.to_thread(calendar_add_event,inputs["summary"],inputs["start"],inputs["end"],inputs.get("description",""),inputs.get("location",""))
    elif name=="drive_search": return await asyncio.to_thread(drive_search_files,inputs["query"],inputs.get("max_results",5))
    elif name=="drive_list": return await asyncio.to_thread(drive_list,inputs.get("folder_id","root"),inputs.get("max_results",20))
    elif name=="drive_read":
        result = await asyncio.to_thread(drive_read_file,inputs["file_id"])
        if isinstance(result,str) and result.startswith("IMAGE_FILE:"):
            # Get the file bytes back and pass to vision
            try:
                svc=build('drive','v3',credentials=get_google_creds())
                meta=svc.files().get(fileId=inputs["file_id"],fields="name").execute()
                filename=meta.get('name','image')
                from googleapiclient.http import MediaIoBaseDownload
                buf=io.BytesIO()
                dl=MediaIoBaseDownload(buf,svc.files().get_media(fileId=inputs["file_id"]))
                done=False
                while not done: _,done=dl.next_chunk()
                return await describe_image_bytes(buf.getvalue(), filename)
            except Exception as e:
                return f"Image vision error: {e}"
        return result
    elif name=="drive_move": return await asyncio.to_thread(drive_move_file,inputs["file_id"],inputs["new_folder_id"])
    elif name=="drive_rename": return await asyncio.to_thread(drive_rename_file,inputs["file_id"],inputs["new_name"])
    elif name=="drive_delete": return await asyncio.to_thread(drive_delete_file,inputs["file_id"])
    elif name=="drive_create_folder": return await asyncio.to_thread(drive_create_folder,inputs["name"],inputs.get("parent_id","root"))
    elif name=="drive_upload": return await asyncio.to_thread(drive_upload_file,inputs["filename"],inputs["content_text"],inputs.get("folder_id","root"))
    elif name=="family_drive_search": return await asyncio.to_thread(drive_search_files,inputs["query"],inputs.get("max_results",5),FAMILY_TOKEN)
    elif name=="family_drive_list": return await asyncio.to_thread(drive_list,inputs.get("folder_id","root"),inputs.get("max_results",20),FAMILY_TOKEN)
    elif name=="family_drive_read": return await asyncio.to_thread(drive_read_file,inputs["file_id"],FAMILY_TOKEN)
    elif name=="contacts_search": return await asyncio.to_thread(contacts_search,inputs["query"],inputs.get("max_results",5))
    elif name=="tasks_list_tasklists": return await asyncio.to_thread(tasks_list_tasklists)
    elif name=="tasks_list": return await asyncio.to_thread(tasks_list,inputs.get("tasklist_id","@default"),inputs.get("show_completed",False))
    elif name=="tasks_add": return await asyncio.to_thread(tasks_add,inputs["title"],inputs.get("notes",""),inputs.get("due"),inputs.get("tasklist_id","@default"))
    elif name=="tasks_complete": return await asyncio.to_thread(tasks_complete,inputs["task_id"],inputs.get("tasklist_id","@default"))
    elif name=="tasks_delete": return await asyncio.to_thread(tasks_delete,inputs["task_id"],inputs.get("tasklist_id","@default"))
    elif name=="outlook_unread": return await asyncio.to_thread(outlook_unread,inputs.get("max_results",10))
    elif name=="outlook_read": return await asyncio.to_thread(outlook_read,inputs["message_id"])
    elif name=="outlook_search": return await asyncio.to_thread(outlook_search,inputs["query"],inputs.get("max_results",10))
    elif name=="outlook_send": return await asyncio.to_thread(outlook_send,inputs["to"],inputs["subject"],inputs["body"])
    elif name=="onenote_notebooks": return await asyncio.to_thread(onenote_list_notebooks)
    elif name=="onenote_sections": return await asyncio.to_thread(onenote_list_sections,inputs.get("notebook_name"))
    elif name=="onenote_list_pages": return await asyncio.to_thread(onenote_list_pages,inputs["section_id"],inputs.get("max_results",20))
    elif name=="onenote_recent": return await asyncio.to_thread(onenote_recent_pages,inputs.get("max_results",10))
    elif name=="onenote_search": return await asyncio.to_thread(onenote_search_pages,inputs["query"],inputs.get("max_results",5))
    elif name=="onenote_read": return await asyncio.to_thread(onenote_get_page,inputs["page_id"])
    elif name=="onenote_create": return await asyncio.to_thread(onenote_create_page,inputs["section_id"],inputs["title"],inputs["content"])
    elif name=="onedrive_search": return await asyncio.to_thread(onedrive_search,inputs["query"],inputs.get("max_results",10))
    elif name=="onedrive_list": return await asyncio.to_thread(onedrive_list,inputs.get("folder_id","root"),inputs.get("max_results",20))
    elif name=="onedrive_read": return await asyncio.to_thread(onedrive_read,inputs["item_id"])
    elif name=="onedrive_upload": return await asyncio.to_thread(onedrive_upload,inputs["filename"],inputs["content_text"],inputs.get("folder_path","/"))
    elif name=="onedrive_delete": return await asyncio.to_thread(onedrive_delete,inputs["item_id"])
    elif name=="plaid_accounts": return await asyncio.to_thread(get_accounts) if PLAID_ENABLED else "Plaid not configured."
    elif name=="plaid_transactions": return await asyncio.to_thread(get_transactions,inputs.get("days",30),inputs.get("max_results",50)) if PLAID_ENABLED else "Plaid not configured."
    elif name=="plaid_spending_by_category": return await asyncio.to_thread(spending_by_category,inputs.get("days",30)) if PLAID_ENABLED else "Plaid not configured."
    elif name=="plaid_debt_snapshot":
        if not PLAID_ENABLED: return "Plaid not configured."
        debts = await asyncio.to_thread(get_debt_snapshot)
        now = datetime.now().strftime("%Y-%m-%d")
        for acct, balance in debts.items():
            memory_save("finances", f"debt_{acct}_{now}", f"${balance:,.2f}")
        memory_save("finances", "last_debt_snapshot", now)
        return f"Debt snapshot saved for {now}: {debts}"
    elif name=="plaid_accounts": return await asyncio.to_thread(get_accounts) if PLAID_ENABLED else "Plaid not configured."
    elif name=="plaid_transactions": return await asyncio.to_thread(get_transactions,inputs.get("days",30),inputs.get("max_results",50)) if PLAID_ENABLED else "Plaid not configured."
    elif name=="plaid_spending_by_category": return await asyncio.to_thread(spending_by_category,inputs.get("days",30)) if PLAID_ENABLED else "Plaid not configured."
    elif name=="plaid_debt_snapshot":
        if not PLAID_ENABLED: return "Plaid not configured."
        debts = await asyncio.to_thread(get_debt_snapshot)
        now = datetime.now().strftime("%Y-%m-%d")
        for acct, balance in debts.items():
            memory_save("finances", f"debt_{acct}_{now}", f"${balance:,.2f}")
        memory_save("finances", "last_debt_snapshot", now)
        return f"Debt snapshot saved for {now}: {debts}"
    elif name=="reminder_set":
        reminder_save(OWNER_TELEGRAM_ID, inputs["message"], inputs["remind_at"])
        return f"Reminder set: '{inputs['message']}' at {inputs['remind_at']} UTC"
    elif name=="reminder_list": return reminder_list(OWNER_TELEGRAM_ID)
    elif name=="reminder_delete": return reminder_delete(inputs["reminder_id"], OWNER_TELEGRAM_ID)
    return f"Unknown tool: {name}"

async def sub_agent(task_description, context=""):
    """Spawn a focused Claude sub-agent to plan and execute a complex task."""
    client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_KEY)
    system = f"""You are a task planning sub-agent for Clawdia. Your job is to break down complex tasks into a clear ordered list of tool calls needed to complete them.

Given a task, respond with ONLY a JSON array of steps. Each step has:
- "description": what this step does
- "tool": the tool name to call
- "inputs": the tool inputs as a dict

Available tools: {json.dumps([t["name"] for t in TOOLS])}

Context about the user: {context}

Respond with ONLY valid JSON, no other text. Example:
[
  {{"description": "List certificates folder", "tool": "family_drive_list", "inputs": {{"folder_id": "root"}}}},
  {{"description": "Read diploma PDF", "tool": "family_drive_read", "inputs": {{"file_id": "abc123"}}}}
]"""

    response = await client.messages.create(
        model=MODEL, max_tokens=2048,
        system=system,
        messages=[{"role": "user", "content": task_description}]
    )
    try:
        plan_text = response.content[0].text.strip()
        # Clean up any markdown
        plan_text = re.sub(r"```json|```", "", plan_text).strip()
        return json.loads(plan_text)
    except Exception as e:
        return None

async def execute_batch_task(chat_id, task_description):
    """Execute a multi-step task using an agentic loop that chains results between steps."""
    client = anthropic.AsyncAnthropic(api_key=ANTHROPIC_KEY)

    # Use a full agentic loop — Claude plans AND executes, chaining results naturally
    system = f"""You are Clawdia executing a background task for Sean. You have full access to all tools.

Execute the task completely and autonomously. When you discover IDs or information in one step, use them in subsequent steps — do not use placeholder values. Work through the entire task until done.

Current memory about Sean:
{memory_load_all()}

When saving to memory, use specific category/key/value — never leave fields empty.
Report what you accomplished when done."""

    messages = [{"role": "user", "content": f"Execute this task completely: {task_description}"}]
    results_log = []

    for _ in range(30):  # up to 30 tool rounds for complex tasks
        response = await client.messages.create(
            model=MODEL, max_tokens=2048,
            system=system, tools=TOOLS, messages=messages
        )

        text_parts = [b.text for b in response.content if b.type == "text"]
        tool_uses = [b for b in response.content if b.type == "tool_use"]

        if not tool_uses:
            # Task complete
            final = "\n".join(text_parts).strip() or "Task completed."
            history_append(chat_id, "user", f"[Task]: {task_description}")
            history_append(chat_id, "assistant", final)
            return final

        messages.append({"role": "assistant", "content": response.content})

        # Execute tools and feed results back
        tool_results = await asyncio.gather(*[run_tool(t.name, t.input) for t in tool_uses])

        for t, result in zip(tool_uses, tool_results):
            results_log.append(f"{t.name}: {str(result)[:200]}")

        messages.append({
            "role": "user",
            "content": [{"type": "tool_result", "tool_use_id": t.id, "content": result}
                       for t, result in zip(tool_uses, tool_results)]
        })

    return f"Task hit step limit. Partial results:\n" + "\n".join(results_log[-10:])

def reminder_save(chat_id, message, remind_at):
    """Save a reminder to the database."""
    now = datetime.utcnow().isoformat()
    with get_conn() as conn:
        conn.execute("INSERT INTO reminders(chat_id,message,remind_at,created) VALUES(?,?,?,?)",
            (chat_id, message, remind_at, now))
    return f"Reminder set: '{message}' at {remind_at}"

def reminder_list(chat_id):
    """List pending reminders."""
    with get_conn() as conn:
        rows = conn.execute("SELECT id,message,remind_at FROM reminders WHERE chat_id=? AND fired=0 ORDER BY remind_at",
            (chat_id,)).fetchall()
    if not rows: return "No pending reminders."
    lines = ["Your pending reminders:"]
    for rid, msg, rat in rows:
        lines.append(f"- [{rid}] {rat[:16]} — {msg}")
    return "\n".join(lines)

def reminder_delete(reminder_id, chat_id):
    """Delete a reminder."""
    with get_conn() as conn:
        n = conn.execute("DELETE FROM reminders WHERE id=? AND chat_id=?", (reminder_id, chat_id)).rowcount
    return "Reminder deleted." if n else "Reminder not found."

async def reminder_scheduler(app, owner_id):
    """Background task that checks and fires reminders every 30 seconds."""
    import asyncio
    log.info("Reminder scheduler started")
    while True:
        try:
            now = datetime.utcnow().isoformat()
            with get_conn() as conn:
                due = conn.execute(
                    "SELECT id,chat_id,message FROM reminders WHERE fired=0 AND remind_at<=? LIMIT 10",
                    (now,)).fetchall()
            for rid, chat_id, message in due:
                try:
                    await app.bot.send_message(chat_id=chat_id, text=f"⏰ Reminder: {message}")
                    with get_conn() as conn:
                        conn.execute("UPDATE reminders SET fired=1 WHERE id=?", (rid,))
                    log.info("Fired reminder %d for chat %d: %s", rid, chat_id, message)
                except Exception as e:
                    log.error("Failed to send reminder %d: %s", rid, e)
        except Exception as e:
            log.error("Reminder scheduler error: %s", e)
        await asyncio.sleep(30)

def build_system_prompt():
    memories=memory_load_all()
    if len(memories)>MAX_MEMORY_CHARS: memories=memories[:MAX_MEMORY_CHARS]+"\n...(truncated)"
    now=datetime.now().strftime("%A, %B %d, %Y %I:%M %p")
    return f"""# Who You Are

You are Clawdia — Sean Durgin's personal AI assistant. Sharp, warm, resourceful. Gets things done without fuss.

Current date/time: {now} (America/New_York)

# Core Truths
Be genuinely helpful, not performatively helpful. Have opinions. Be resourceful before asking.
Earn trust through competence. NEVER send email without explicit confirmation. Private things stay private.

# About Sean
- Sean Durgin, North East MD (home) / Northern Virginia (work)
- Retired USAF Master Sergeant, 21+ years, Cyber Defense Operations, discharged Feb 1 2024
- Job: Data center technician at Oracle
- Email: seandurgin@gmail.com (personal), durginfamily@gmail.com (family)
- Notes: OneNote preferred

# Your Persistent Memory
{memories}

# Your Tools (40 total)
Google (seandurgin): gmail_unread, gmail_read, gmail_send, calendar_upcoming, calendar_add, drive_search, drive_list, drive_read, drive_move, drive_rename, drive_delete, drive_create_folder, drive_upload, contacts_search
Google (durginfamily): family_gmail_unread, family_gmail_read, family_gmail_send, family_drive_search, family_drive_list, family_drive_read
Microsoft: onenote_notebooks, onenote_sections, onenote_list_pages, onenote_recent, onenote_search, onenote_read, onenote_create, onedrive_search, onedrive_list, onedrive_read, onedrive_upload, onedrive_delete
Web: web_search, fetch_url
Reminders: reminder_set, reminder_list, reminder_delete
Finance (on-demand only): plaid_accounts, plaid_transactions, plaid_spending_by_category, plaid_debt_snapshot
Finance (on-demand only): plaid_accounts, plaid_transactions, plaid_spending_by_category, plaid_debt_snapshot
Memory: save_memory, delete_memory

When Sean uploads a file or image, the contents are provided directly — no tool needed.

# Memory Discipline
Save facts about Sean immediately when he shares them. Your memory is how you persist.
"""

async def ask_claude(chat_id, user_text):
    client=anthropic.AsyncAnthropic(api_key=ANTHROPIC_KEY)
    history_append(chat_id,"user",user_text)
    messages=history_get(chat_id)
    system=build_system_prompt()
    for _ in range(30):
        response=await client.messages.create(model=MODEL,max_tokens=1024,system=system,tools=TOOLS,messages=messages)
        text_parts=[b.text for b in response.content if b.type=="text"]
        tool_uses=[b for b in response.content if b.type=="tool_use"]
        if not tool_uses:
            final_text="\n".join(text_parts).strip() or "(no response)"
            history_append(chat_id,"assistant",final_text)
            return final_text
        messages.append({"role":"assistant","content":response.content})
        tool_results=await asyncio.gather(*[run_tool(t.name,t.input) for t in tool_uses])
        messages.append({"role":"user","content":[{"type":"tool_result","tool_use_id":t.id,"content":result} for t,result in zip(tool_uses,tool_results)]})
    return "I got stuck. Could you rephrase?"

def is_authorized(update):
    return OWNER_TELEGRAM_ID==0 or update.effective_user.id==OWNER_TELEGRAM_ID

async def handle_message(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not update.message or not update.message.text or not is_authorized(update): return
    chat_id=update.effective_chat.id; user_msg=update.message.text.strip()
    log.info("User [%s]: %s",chat_id,user_msg[:80])
    await context.bot.send_chat_action(chat_id=chat_id,action=ChatAction.TYPING)
    try: reply=await ask_claude(chat_id,user_msg)
    except Exception as e: log.exception("Error"); reply=f"Something went wrong: {e}"
    await update.message.reply_text(reply)

async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_authorized(update): return
    doc=update.message.document
    if not doc: return
    chat_id=update.effective_chat.id
    filename=doc.file_name or "unknown_file"
    log.info("File received [%s]: %s",chat_id,filename)
    await context.bot.send_chat_action(chat_id=chat_id,action=ChatAction.TYPING)
    try:
        tfile=await context.bot.get_file(doc.file_id)
        file_bytes=bytes(await tfile.download_as_bytearray())
        text=extract_file_text(file_bytes,filename)
        caption=update.message.caption or ""
        prompt=f"I uploaded a file called '{filename}'. Contents:\n\n{text}"
        if caption: prompt+=f"\n\nMy note: {caption}"
        else: prompt+="\n\nPlease summarize this and highlight anything important."
        reply=await ask_claude(chat_id,prompt)
        await update.message.reply_text(reply)
    except Exception as e:
        log.exception("File error")
        await update.message.reply_text(f"Error reading file: {e}")

async def handle_image(update: Update, context: ContextTypes.DEFAULT_TYPE):
    if not is_authorized(update): return
    photo=update.message.photo
    if not photo: return
    chat_id=update.effective_chat.id
    caption=update.message.caption or ""
    log.info("Image received [%s]",chat_id)
    await context.bot.send_chat_action(chat_id=chat_id,action=ChatAction.TYPING)
    try:
        file=await context.bot.get_file(photo[-1].file_id)
        file_bytes=bytes(await file.download_as_bytearray())
        img_b64=base64.b64encode(file_bytes).decode('utf-8')
        client=anthropic.AsyncAnthropic(api_key=ANTHROPIC_KEY)
        question=caption if caption else "Describe this image in detail. If there is any text, read all of it. If it is a document, summarize it."
        response=await client.messages.create(
            model=MODEL,max_tokens=1024,system=build_system_prompt(),
            messages=[{"role":"user","content":[
                {"type":"image","source":{"type":"base64","media_type":"image/jpeg","data":img_b64}},
                {"type":"text","text":question}
            ]}]
        )
        reply=response.content[0].text
        history_append(chat_id,"user",f"[Image]{': '+caption if caption else ''}")
        history_append(chat_id,"assistant",reply)
        await update.message.reply_text(reply)
    except Exception as e:
        log.exception("Image error")
        await update.message.reply_text(f"Error reading image: {e}")

async def cmd_task(update, context: ContextTypes.DEFAULT_TYPE):
    """Execute a long multi-step task silently."""
    if not is_authorized(update): return
    chat_id = update.effective_chat.id
    task = " ".join(context.args) if context.args else ""
    if not task:
        await update.message.reply_text("Usage: /task <describe what you want done>\nExample: /task scan all PDFs in my certificates folder and save what each one is to memory")
        return
    await update.message.reply_text("On it. I'll let you know when done.")
    await context.bot.send_chat_action(chat_id=chat_id, action=ChatAction.TYPING)
    try:
        result = await execute_batch_task(chat_id, task)
        await update.message.reply_text(result)
    except Exception as e:
        log.exception("Batch task error")
        await update.message.reply_text(f"Task failed: {e}")

async def cmd_start(update,context):
    if not is_authorized(update): return
    await update.message.reply_text("Hey Sean — I'm back. What's up?")

async def cmd_memory(update,context):
    if not is_authorized(update): return
    await update.message.reply_text(f"Here's what I remember:\n\n{memory_load_all()}")

async def cmd_forget(update,context):
    if not is_authorized(update): return
    args=context.args
    if len(args)<2: await update.message.reply_text("Usage: /forget <category> <key>"); return
    await update.message.reply_text("Deleted." if memory_delete(args[0]," ".join(args[1:])) else "Not found.")

async def cmd_clearhistory(update,context):
    if not is_authorized(update): return
    with get_conn() as conn: conn.execute("DELETE FROM history WHERE chat_id=?",(update.effective_chat.id,))
    await update.message.reply_text("Conversation history cleared. Memories intact.")

async def cmd_help(update,context):
    if not is_authorized(update): return
    await update.message.reply_text("*Clawdia Commands*\n\n/memory — what I remember\n/forget <category> <key> — delete a memory\n/clearhistory — clear recent chat\n/help — this\n\nSend files or photos and I'll read them.",parse_mode="Markdown")

def main():
    init_db()
    log.info("Starting Clawdia (model: %s, tools: %d)",MODEL,len(TOOLS))
    app=Application.builder().token(TELEGRAM_TOKEN).build()
    from briefing import start_briefing_scheduler
    start_briefing_scheduler(app,OWNER_TELEGRAM_ID,gmail_get_unread,calendar_get_upcoming,brave_search)
    app.add_handler(CommandHandler("task",cmd_task))
    app.add_handler(CommandHandler("start",cmd_start))
    app.add_handler(CommandHandler("memory",cmd_memory))
    app.add_handler(CommandHandler("forget",cmd_forget))
    app.add_handler(CommandHandler("clearhistory",cmd_clearhistory))
    app.add_handler(CommandHandler("help",cmd_help))
    app.add_handler(MessageHandler(filters.Document.ALL,handle_file))
    app.add_handler(MessageHandler(filters.PHOTO,handle_image))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND,handle_message))
    log.info("Clawdia is online.")
    import asyncio
    async def post_init(app):
        asyncio.create_task(reminder_scheduler(app, OWNER_TELEGRAM_ID))
    app.post_init = post_init
    app.run_polling(drop_pending_updates=True)

if __name__=="__main__":
    main()
