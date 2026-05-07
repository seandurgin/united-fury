"""Create Google Docs and .docx files in Drive.

Two formats:
  - "gdoc" -> native Google Doc (uses Drive API to create with HTML body
             converted to Doc format). Returns a shareable Google Doc URL.
  - "docx" -> real Microsoft Word file (built with python-docx, uploaded
             to Drive). Returns a Drive URL where Sean can download the
             actual .docx file. Better for WGU submissions.

Permissions: defaults to anyone-with-link-can-edit, same as create_google_sheet.
"""
import io, logging, re
from googleapiclient.http import MediaIoBaseUpload

log = logging.getLogger("clawdia.google_docs")

GDOC_MIME = "application/vnd.google-apps.document"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _build_docx_bytes(content):
    """Build a .docx in memory from markdown-ish content."""
    import docx
    from docx.shared import Pt
    doc = docx.Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in content.split("\n"):
        line = line.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.lstrip().startswith("- ") or line.lstrip().startswith("* "):
            text = line.lstrip()[2:].strip()
            doc.add_paragraph(_strip_md(text), style="List Bullet")
        elif len(line) > 2 and line[0].isdigit() and line[1:].lstrip().startswith(". "):
            text = line.split(".", 1)[1].lstrip()
            doc.add_paragraph(_strip_md(text), style="List Number")
        else:
            p = doc.add_paragraph()
            _add_runs_with_bold(p, line)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def _strip_md(text):
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def _add_runs_with_bold(paragraph, text):
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _content_to_html(title, content):
    """Convert markdown-ish content to HTML for Drive's Google Doc importer."""
    out = ["<html><head><title>", _html_escape(title), "</title></head><body>"]
    para_buf = []

    def flush():
        if para_buf:
            joined = " ".join(para_buf)
            joined = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", joined)
            out.append("<p>")
            out.append(joined)
            out.append("</p>")
            para_buf.clear()

    for line in content.split("\n"):
        s = line.strip()
        if not s:
            flush()
            continue
        if s.startswith("### "):
            flush()
            out.append("<h3>" + _html_escape(s[4:]) + "</h3>")
        elif s.startswith("## "):
            flush()
            out.append("<h2>" + _html_escape(s[3:]) + "</h2>")
        elif s.startswith("# "):
            flush()
            out.append("<h1>" + _html_escape(s[2:]) + "</h1>")
        elif s.startswith("- ") or s.startswith("* "):
            flush()
            out.append("<ul><li>" + _html_escape(s[2:]) + "</li></ul>")
        else:
            para_buf.append(_html_escape(s))
    flush()
    out.append("</body></html>")
    return "".join(out)


def _html_escape(s):
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def create_google_doc(title, content, format, get_creds_fn, share_anyone=True):
    """Create a Google Doc or .docx file in personal Drive.
    Returns a string with the view/download URL.
    format: 'gdoc' for native Google Doc, 'docx' for real Word file.
    """
    if not title or not isinstance(title, str):
        return "ERROR: create_google_doc requires a non-empty title."
    if not content or not isinstance(content, str):
        return "ERROR: create_google_doc requires non-empty content."
    fmt = (format or "docx").lower()
    if fmt not in ("gdoc", "docx"):
        return "ERROR: format must be 'gdoc' or 'docx', got " + repr(format) + "."

    try:
        from googleapiclient.discovery import build
        creds = get_creds_fn()
        svc = build("drive", "v3", credentials=creds)

        if fmt == "docx":
            docx_bytes = _build_docx_bytes(content)
            media = MediaIoBaseUpload(docx_bytes, mimetype=DOCX_MIME, resumable=False)
            file_metadata = {"name": title if title.endswith(".docx") else (title + ".docx")}
            f = svc.files().create(body=file_metadata, media_body=media,
                                   fields="id,name,webViewLink").execute()
        else:
            html_body = _content_to_html(title, content)
            media = MediaIoBaseUpload(io.BytesIO(html_body.encode("utf-8")),
                                      mimetype="text/html", resumable=False)
            file_metadata = {"name": title, "mimeType": GDOC_MIME}
            f = svc.files().create(body=file_metadata, media_body=media,
                                   fields="id,name,webViewLink").execute()

        if share_anyone:
            try:
                svc.permissions().create(
                    fileId=f["id"],
                    body={"type": "anyone", "role": "writer"},
                    fields="id"
                ).execute()
            except Exception as pe:
                log.warning("permission set failed: %s", pe)

        kind = "Google Doc" if fmt == "gdoc" else ".docx file in Drive"
        return ("Created " + kind + ": " + f.get("name", "?") + "\n" +
                f.get("webViewLink", "(no link)") + "\n\nFile ID: " + f.get("id", "?"))
    except Exception as e:
        return "create_google_doc error: " + str(e)
